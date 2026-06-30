from pathlib import Path

from docx import Document


FILES = [
    r"D:\科研\雄安专项\外场\合-测试报告.docx",
    r"D:\科研\雄安专项\北京交通大学测试报告250619.docx",
    r"D:\科研\雄安专项\修改——最终科技报告-智慧城市立体通信下跨网协同可靠传输技术.docx",
    r"D:\科研\多模态\多模态第三方测试报告1023.docx",
    r"D:\科研\多模态\项目综合绩效自评价.docx",
]


def extract_docx(path: str) -> str:
    doc = Document(path)
    texts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            texts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                text = " / ".join(
                    part.strip() for part in cell.text.splitlines() if part.strip()
                )
                if text:
                    cells.append(text)
            if cells:
                texts.append(" | ".join(cells))
    return "\n".join(texts)


def main() -> None:
    out_dir = Path("doc_extracts")
    out_dir.mkdir(exist_ok=True)
    for raw in FILES:
        path = Path(raw)
        text = extract_docx(raw)
        out_path = out_dir / f"{path.stem}.txt"
        out_path.write_text(text, encoding="utf-8")
        print(f"--- {path.name} chars={len(text)} -> {out_path}")
        print(text[:1800])
        print()


if __name__ == "__main__":
    main()
